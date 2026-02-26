"""Enhanced DOCX extraction strategy for RAG.

Extracts:
- Hierarchical structure (via paragraph styles)
- Tables in Markdown format
- Metadata (paragraph count, heading count)
"""

import logging

from docx import Document

from .base import BaseExtractionStrategy, ExtractionResult, Page, Section, Table

logger = logging.getLogger(__name__)


class DocxStructuredStrategy(BaseExtractionStrategy):
    """RAG-optimized DOCX extraction with heading and table support."""

    def extract(self, file_stream, language: str = "eng") -> ExtractionResult:
        """Extract text, tables, and hierarchy from DOCX."""

        logger.info("Starting structured Docx extraction.")
        try:
            if hasattr(file_stream, "seek"):
                file_stream.seek(0)

            # Check file size warning
            if hasattr(file_stream, "seek") and hasattr(file_stream, "tell"):
                file_stream.seek(0, 2)
                file_size_mb = file_stream.tell() / (1024 * 1024)
                file_stream.seek(0)

                if file_size_mb > 50:
                    logger.warning(f"Large DOCX file: {file_size_mb:.1f} MB")

            doc = Document(file_stream)

            # Extract hierarchy from paragraph styles
            hierarchy = []
            full_text_parts = []
            current_section = None
            section_content = []

            for para in doc.paragraphs:
                text = para.text.strip()

                if not text:
                    continue

                # Detect headings
                style_name = para.style.name if para.style else "Normal"

                if "Heading" in style_name:
                    # Save previous section
                    if current_section:
                        current_section.content = "\n".join(section_content)
                        hierarchy.append(current_section)

                    # Extract heading level
                    level = self._extract_heading_level(style_name)

                    # Start new section
                    current_section = Section(
                        level=level,
                        title=text,
                        content="",
                        start_page=1,  # DOCX doesn't expose page breaks easily
                        metadata={"style": style_name},
                    )
                    section_content = []

                    # Add heading to full text with markdown
                    heading_prefix = "#" * level
                    full_text_parts.append(f"{heading_prefix} {text}")
                else:
                    # Regular paragraph
                    full_text_parts.append(text)
                    if current_section:
                        section_content.append(text)

            # Save last section
            if current_section:
                current_section.content = "\n".join(section_content)
                hierarchy.append(current_section)

            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                markdown_table = self._table_to_markdown(table)
                tables.append(
                    Table(
                        markdown=markdown_table,
                        caption=None,
                        page_number=1,
                        metadata={"table_index": table_idx},
                    )
                )
                full_text_parts.append(f"\n{markdown_table}\n")

            full_text = "\n\n".join(full_text_parts)

            # Create single page (DOCX doesn't expose pages)
            pages = [
                Page(
                    page_number=1,
                    content=full_text,
                    metadata={
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(tables),
                    },
                )
            ]

            logger.info(
                f"Structured DOCX extraction complete: "
                f"{len(hierarchy)} sections, {len(tables)} tables"
            )

            return ExtractionResult(
                full_text=full_text,
                pages=pages,
                hierarchy=hierarchy,
                tables=tables,
                metadata={
                    "method": "docx_structured",
                    "page_count": 1,
                    "paragraph_count": len(doc.paragraphs),
                    "section_count": len(hierarchy),
                    "table_count": len(tables),
                },
            )

        except Exception as e:
            logger.error(f"Docx extraction failed: {e}")
            raise Exception(f"Failed to extract text from Docx: {str(e)}")

    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name.

        Args:
            style_name: e.g., "Heading 1", "Heading 2"

        Returns:
            Heading level (1-3, capped)
        """
        import re

        match = re.search(r"Heading (\d+)", style_name)
        if match:
            return min(int(match.group(1)), 3)
        return 1

    def _table_to_markdown(self, table) -> str:
        """Convert DOCX table to Markdown.

        Args:
            table: python-docx Table object

        Returns:
            Markdown table string
        """
        if not table.rows:
            return ""

        lines = []

        # Header row (first row)
        header_cells = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header_cells) + " |")

        # Separator
        lines.append("|" + "|".join("---" for _ in header_cells) + "|")

        # Data rows
        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")

        return "\n".join(lines)
